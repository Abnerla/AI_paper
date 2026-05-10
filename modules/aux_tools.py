# -*- coding: utf-8 -*-
"""
公共文档处理工具 - Word导入导出、文本统计、差异对比
"""

import os
import re
import sys
import difflib
import tempfile
from collections import Counter

from modules.table_blocks import (
    TABLE_ALIGN_CENTER,
    TABLE_ALIGN_RIGHT,
    TABLE_STYLE_GRID,
    TABLE_STYLE_THREE_LINE,
    blocks_to_plain_text,
    new_paragraph_block,
    new_table_block,
    normalize_table_alignments,
    normalize_table_alignment,
    normalize_merged_cells,
    parse_markdown_blocks,
    sanitize_blocks,
)


class AuxTools:
    """公共文档处理能力集合"""

    def __init__(self, api_client=None):
        self.api = api_client

    _LATEX_SPECIAL_CHARS = {
        '\\': r'\textbackslash{}',
        '{': r'\{',
        '}': r'\}',
        '$': r'\$',
        '&': r'\&',
        '#': r'\#',
        '%': r'\%',
        '_': r'\_',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }

    def _escape_latex(self, text: str) -> str:
        """转义 LaTeX 保留字符。"""
        return ''.join(self._LATEX_SPECIAL_CHARS.get(char, char) for char in str(text or ''))

    def _latex_heading_command(self, line: str) -> str:
        """将常见标题行映射为 LaTeX 节命令。"""
        stripped = str(line or '').strip()
        if not stripped:
            return ''

        markdown_match = re.match(r'^(#{1,3})\s+(.*)$', stripped)
        if markdown_match:
            level = len(markdown_match.group(1))
            title = self._escape_latex(markdown_match.group(2).strip())
            if level == 1:
                return rf'\section{{{title}}}'
            if level == 2:
                return rf'\subsection{{{title}}}'
            return rf'\subsubsection{{{title}}}'

        if re.match(r'^第[一二三四五六七八九十百千万\d]+[章节部分篇]\s*', stripped):
            return rf'\section{{{self._escape_latex(stripped)}}}'

        if re.match(r'^[一二三四五六七八九十百千万\d]+[、.．)]\s*', stripped):
            return rf'\subsection{{{self._escape_latex(stripped)}}}'

        if re.match(r'^\d+(?:\.\d+){0,2}\s+', stripped):
            level = stripped.count('.')
            command = 'section' if level == 0 else 'subsection' if level == 1 else 'subsubsection'
            return rf'\{command}{{{self._escape_latex(stripped)}}}'

        return ''

    def _latex_body_from_text(self, text: str) -> str:
        """将普通文本整理为 LaTeX 正文。"""
        lines = []
        paragraph_buffer = []

        def flush_paragraph():
            if not paragraph_buffer:
                return
            paragraph_text = '\n'.join(part.strip() for part in paragraph_buffer if part.strip())
            paragraph_buffer.clear()
            if paragraph_text:
                lines.append(self._escape_latex(paragraph_text))
                lines.append('')

        for raw_line in str(text or '').splitlines():
            stripped = raw_line.strip()
            if not stripped:
                flush_paragraph()
                continue

            heading_command = self._latex_heading_command(stripped)
            if heading_command:
                flush_paragraph()
                lines.append(heading_command)
                lines.append('')
                continue

            paragraph_buffer.append(raw_line.rstrip())

        flush_paragraph()
        return '\n'.join(lines).rstrip()

    @staticmethod
    def _looks_like_table_caption(text: str) -> bool:
        normalized = re.sub(r'\s+', ' ', str(text or '').strip())
        return bool(
            re.match(
                r'^(?:表|Table)\s*[\d一二三四五六七八九十百千万IVXLCDMivxlcdm]+(?:[\s.．、:：-]|$)',
                normalized,
            )
        )

    @staticmethod
    def _docx_style_name(style) -> str:
        try:
            return str(getattr(style, 'name', '') or '').strip()
        except Exception:
            return ''

    @staticmethod
    def _docx_style_id(style) -> str:
        try:
            return str(getattr(style, 'style_id', '') or '').strip()
        except Exception:
            return ''

    @staticmethod
    def _docx_outline_level_from_element(element) -> int:
        try:
            from docx.oxml.ns import qn
            p_pr = element.pPr
            if p_pr is None:
                return -1
            outline = p_pr.find(qn('w:outlineLvl'))
            if outline is None:
                return -1
            value = outline.get(qn('w:val'))
            if value is None:
                return -1
            return max(0, int(value))
        except Exception:
            return -1

    @classmethod
    def _docx_paragraph_outline_level(cls, paragraph) -> int:
        direct_level = cls._docx_outline_level_from_element(paragraph._p)
        if direct_level >= 0:
            return direct_level

        try:
            style_element = paragraph.style.element
        except Exception:
            style_element = None
        if style_element is None:
            return -1
        return cls._docx_outline_level_from_element(style_element)

    @staticmethod
    def _docx_run_font_size_pt(run):
        try:
            size = run.font.size
            if size is not None:
                return float(size.pt)
        except Exception:
            pass
        try:
            style_size = run.style.font.size
            if style_size is not None:
                return float(style_size.pt)
        except Exception:
            pass
        return None

    @classmethod
    def _docx_paragraph_font_size_pt(cls, paragraph):
        sizes = []
        for run in paragraph.runs:
            text = str(run.text or '')
            if not text.strip():
                continue
            size = cls._docx_run_font_size_pt(run)
            if size is not None:
                sizes.append(round(size, 2))
        if sizes:
            return Counter(sizes).most_common(1)[0][0]
        try:
            style_size = paragraph.style.font.size
            if style_size is not None:
                return round(float(style_size.pt), 2)
        except Exception:
            pass
        return None

    @staticmethod
    def _docx_run_is_bold(run):
        try:
            if run.bold is not None:
                return bool(run.bold)
        except Exception:
            pass
        try:
            if run.style.font.bold is not None:
                return bool(run.style.font.bold)
        except Exception:
            pass
        return False

    @classmethod
    def _docx_paragraph_bold_ratio(cls, paragraph) -> float:
        total = 0
        bold = 0
        for run in paragraph.runs:
            text = str(run.text or '')
            visible_len = len(text.strip())
            if visible_len <= 0:
                continue
            total += visible_len
            if cls._docx_run_is_bold(run):
                bold += visible_len
        if total <= 0:
            try:
                return 1.0 if bool(paragraph.style.font.bold) else 0.0
            except Exception:
                return 0.0
        return bold / total

    @staticmethod
    def _docx_paragraph_alignment(paragraph) -> str:
        value = paragraph.alignment
        if value is None:
            return ''
        name = getattr(value, 'name', '')
        if name:
            return str(name).lower()
        try:
            return str(int(value))
        except Exception:
            return str(value)

    @classmethod
    def _docx_cell_alignment(cls, tc, parent_table) -> str:
        try:
            from docx.table import _Cell
            cell = _Cell(tc, parent_table)
            for paragraph in cell.paragraphs:
                alignment = cls._docx_paragraph_alignment(paragraph)
                if alignment:
                    return normalize_table_alignment(alignment)
        except Exception:
            pass
        return normalize_table_alignment('')

    @staticmethod
    def _looks_like_toc_paragraph(text: str, style_name: str = '', style_id: str = '') -> bool:
        normalized = re.sub(r'\s+', ' ', str(text or '').strip())
        plain = normalized.strip('：:').lower()
        style_hint = f'{style_name} {style_id}'.lower()
        if not normalized:
            return False
        if 'toc' in style_hint or '目录' in style_hint:
            return True
        if plain in {'目录', 'contents', 'table of contents'}:
            return True
        if re.search(r'(?:\.{2,}|…{2,}|·{2,}|_{2,})\s*\d+\s*$', normalized):
            return True
        return bool(re.search(r'\s+\d+\s*$', normalized) and re.match(r'^(?:第[一二三四五六七八九十百千万\d]+章|[一二三四五六七八九十百千万\d]+(?:\.\d+){0,3})\s+', normalized))

    @classmethod
    def _docx_paragraph_block(cls, paragraph, text: str):
        style = getattr(paragraph, 'style', None)
        style_name = cls._docx_style_name(style)
        style_id = cls._docx_style_id(style)
        metadata = {
            'style_name': style_name,
            'style_id': style_id,
            'outline_level': cls._docx_paragraph_outline_level(paragraph),
            'font_size_pt': cls._docx_paragraph_font_size_pt(paragraph),
            'bold_ratio': cls._docx_paragraph_bold_ratio(paragraph),
            'alignment': cls._docx_paragraph_alignment(paragraph),
            'is_toc_like': cls._looks_like_toc_paragraph(text, style_name, style_id),
        }
        return new_paragraph_block(text, **metadata)

    @staticmethod
    def _docx_grid_span(tc) -> int:
        try:
            grid_span = tc.tcPr.gridSpan if tc.tcPr is not None else None
            return max(1, int(grid_span.val)) if grid_span is not None else 1
        except Exception:
            return 1

    @staticmethod
    def _docx_vmerge_value(tc) -> str:
        try:
            vmerge = tc.tcPr.vMerge if tc.tcPr is not None else None
            if vmerge is None:
                return ''
            value = vmerge.val
            return str(value or 'continue').strip().lower()
        except Exception:
            return ''

    @staticmethod
    def _extract_docx_cell_text(tc, parent_table) -> str:
        try:
            from docx.table import _Cell
            return str(_Cell(tc, parent_table).text or '').replace('\r', '').strip()
        except Exception:
            try:
                paragraphs = []
                for paragraph in tc.xpath('./w:p'):
                    text = ''.join(paragraph.xpath('.//w:t/text()')).strip()
                    if text:
                        paragraphs.append(text)
                return '\n'.join(paragraphs).strip()
            except Exception:
                try:
                    return ''.join(tc.xpath('.//w:t/text()')).strip()
                except Exception:
                    return ''

    @classmethod
    def _docx_table_column_count(cls, table) -> int:
        grid_count = 0
        try:
            grid_count = len(table._tbl.tblGrid.gridCol_lst)
        except Exception:
            grid_count = 0

        row_width = 0
        try:
            for tr in table._tbl.tr_lst:
                width = 0
                for tc in tr.tc_lst:
                    width += cls._docx_grid_span(tc)
                row_width = max(row_width, width)
        except Exception:
            row_width = 0
        return max(1, grid_count, row_width)

    @staticmethod
    def _docx_cell_border_value(tc, edge: str) -> str:
        try:
            from docx.oxml.ns import qn
            tc_pr = tc.tcPr
            tc_borders = tc_pr.first_child_found_in('w:tcBorders') if tc_pr is not None else None
            if tc_borders is None:
                return ''
            element = tc_borders.find(qn(f'w:{edge}'))
            if element is None:
                return ''
            return str(element.get(qn('w:val')) or 'single').strip().lower()
        except Exception:
            return ''

    @staticmethod
    def _docx_table_border_value(table, edge: str) -> str:
        try:
            from docx.oxml.ns import qn
            tbl_pr = table._tbl.tblPr
            tbl_borders = tbl_pr.first_child_found_in('w:tblBorders') if tbl_pr is not None else None
            if tbl_borders is None:
                return ''
            element = tbl_borders.find(qn(f'w:{edge}'))
            if element is None:
                return ''
            return str(element.get(qn('w:val')) or 'single').strip().lower()
        except Exception:
            return ''

    @classmethod
    def _detect_docx_table_style(cls, table) -> str:
        try:
            style_name = str(getattr(getattr(table, 'style', None), 'name', '') or '').lower()
            if '三线' in style_name or ('three' in style_name and 'line' in style_name):
                return TABLE_STYLE_THREE_LINE
        except Exception:
            pass

        no_line_values = {'nil', 'none'}

        def has_line(value):
            return bool(value) and value not in no_line_values

        def no_line(value):
            return bool(value) and value in no_line_values

        table_top = cls._docx_table_border_value(table, 'top')
        table_bottom = cls._docx_table_border_value(table, 'bottom')
        table_vertical = [
            cls._docx_table_border_value(table, edge)
            for edge in ('left', 'right', 'insideV')
        ]
        if has_line(table_top) and has_line(table_bottom) and not any(has_line(value) for value in table_vertical):
            return TABLE_STYLE_THREE_LINE

        try:
            rows = list(table._tbl.tr_lst)
        except Exception:
            rows = []
        if not rows:
            return TABLE_STYLE_GRID

        first_row = list(rows[0].tc_lst)
        last_row = list(rows[-1].tc_lst)
        if not first_row or not last_row:
            return TABLE_STYLE_GRID

        def all_cells_have_line(cells, edge):
            values = [cls._docx_cell_border_value(tc, edge) for tc in cells]
            values = [value for value in values if value]
            return bool(values) and all(has_line(value) for value in values)

        vertical_values = []
        for tr in rows:
            for tc in tr.tc_lst:
                vertical_values.append(cls._docx_cell_border_value(tc, 'left'))
                vertical_values.append(cls._docx_cell_border_value(tc, 'right'))
        vertical_values = [value for value in vertical_values if value]
        vertical_no_line = bool(vertical_values) and all(no_line(value) for value in vertical_values)

        if (
            vertical_no_line
            and all_cells_have_line(first_row, 'top')
            and all_cells_have_line(first_row, 'bottom')
            and all_cells_have_line(last_row, 'bottom')
        ):
            return TABLE_STYLE_THREE_LINE
        return TABLE_STYLE_GRID

    @classmethod
    def _docx_table_to_block(cls, table, caption: str = ''):
        col_count = cls._docx_table_column_count(table)
        rows = []
        cell_alignments = []
        merged_cells = []
        active_vmerges = {}

        for row_idx, tr in enumerate(table._tbl.tr_lst):
            row_values = [''] * col_count
            row_alignments = [normalize_table_alignment('') for _col in range(col_count)]
            col_idx = 0
            next_vmerges = {}
            for tc in tr.tc_lst:
                if col_idx >= col_count:
                    break
                colspan = min(cls._docx_grid_span(tc), col_count - col_idx)
                vmerge = cls._docx_vmerge_value(tc)
                if vmerge == 'continue':
                    active = active_vmerges.get(col_idx)
                    if active is not None:
                        active['rowspan'] += 1
                        next_vmerges[col_idx] = active
                else:
                    row_values[col_idx] = cls._extract_docx_cell_text(tc, table)
                    alignment = cls._docx_cell_alignment(tc, table)
                    for span_idx in range(col_idx, min(col_idx + colspan, col_count)):
                        row_alignments[span_idx] = alignment
                    if vmerge == 'restart':
                        active = {
                            'row': row_idx,
                            'col': col_idx,
                            'rowspan': 1,
                            'colspan': colspan,
                        }
                        merged_cells.append(active)
                        next_vmerges[col_idx] = active
                    elif colspan > 1:
                        merged_cells.append({
                            'row': row_idx,
                            'col': col_idx,
                            'rowspan': 1,
                            'colspan': colspan,
                        })
                col_idx += colspan
            rows.append(row_values)
            cell_alignments.append(row_alignments)
            active_vmerges = next_vmerges

        return new_table_block(
            rows or [['']],
            caption=caption,
            merged_cells=merged_cells,
            table_style=cls._detect_docx_table_style(table),
            cell_alignments=cell_alignments,
        )

    def import_docx_blocks(self, filepath: str) -> dict:
        """导入 Word 文档，返回结构化内容块。"""
        try:
            normalized_path = os.path.abspath(os.path.expanduser(str(filepath or '').strip()))
            if not normalized_path.lower().endswith('.docx'):
                raise RuntimeError('仅支持导入 .docx 格式的 Word 文档')
            import docx
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            doc = docx.Document(normalized_path)
            blocks = []
            for child in doc.element.body.iterchildren():
                if isinstance(child, CT_P):
                    paragraph = Paragraph(child, doc)
                    text = str(paragraph.text or '').strip()
                    if text:
                        blocks.append(self._docx_paragraph_block(paragraph, text))
                    continue
                if isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    caption = ''
                    if blocks and blocks[-1].get('type') == 'paragraph':
                        previous_text = str(blocks[-1].get('text', '') or '').strip()
                        if self._looks_like_table_caption(previous_text):
                            caption = previous_text
                            blocks.pop()
                    blocks.append(self._docx_table_to_block(table, caption=caption))
            sanitized = sanitize_blocks(blocks)
            return {
                'text': blocks_to_plain_text(sanitized),
                'blocks': sanitized,
            }
        except ImportError:
            raise RuntimeError('请安装python-docx库: pip install python-docx')
        except Exception as e:
            raise RuntimeError(f'读取Word文件失败: {e}')

    def import_docx(self, filepath: str) -> str:
        """导入 Word 文档，返回文本内容。"""
        return self.import_docx_blocks(filepath).get('text', '')

    def export_docx(self, text: str, filepath: str, title: str = '', level_font_styles: dict = None, sections_data: dict = None) -> bool:
        """导出为Word文档。sections_data 可包含 section_order, sections, section_levels, section_blocks。"""
        try:
            import docx
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx.Document()

            # 设置页边距
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)

            lfs = level_font_styles or {}
            body_style = lfs.get('body', {})
            h1_style = lfs.get('h1', {})
            h2_style = lfs.get('h2', {})
            h3_style = lfs.get('h3', {})

            body_font = body_style.get('font', '宋体')
            body_font_en = body_style.get('font_en', 'Times New Roman')
            body_pt = body_style.get('size_pt', 12)
            h1_font = h1_style.get('font', '黑体')
            h1_font_en = h1_style.get('font_en', 'Times New Roman')
            h1_pt = h1_style.get('size_pt', 16)
            h2_font = h2_style.get('font', '黑体')
            h2_font_en = h2_style.get('font_en', 'Times New Roman')
            h2_pt = h2_style.get('size_pt', 14)
            h3_font = h3_style.get('font', '黑体')
            h3_font_en = h3_style.get('font_en', 'Times New Roman')
            h3_pt = h3_style.get('size_pt', 12)

            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def _set_run_font(run, cn_font, en_font, pt_size):
                """设置 run 的中英文字体和字号"""
                run.font.name = en_font
                run.font.size = Pt(pt_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)

            def _set_cell_borders(cell, **borders):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.first_child_found_in('w:tcBorders')
                if tc_borders is None:
                    tc_borders = OxmlElement('w:tcBorders')
                    tc_pr.append(tc_borders)

                for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    if edge not in borders:
                        continue
                    tag = f'w:{edge}'
                    element = tc_borders.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tc_borders.append(element)
                    config = borders[edge] or {}
                    if not config:
                        element.set(qn('w:val'), 'nil')
                        continue
                    for key, value in config.items():
                        element.set(qn(f'w:{key}'), str(value))

            def _apply_three_line_table_style(table):
                row_count = len(table.rows)
                col_count = len(table.columns)
                if row_count <= 0 or col_count <= 0:
                    return
                no_border = {'val': 'nil'}
                strong = {'val': 'single', 'sz': '12', 'space': '0', 'color': '000000'}
                normal = {'val': 'single', 'sz': '8', 'space': '0', 'color': '000000'}
                for row_idx in range(row_count):
                    for col_idx in range(col_count):
                        borders = {
                            'top': no_border,
                            'left': no_border,
                            'bottom': no_border,
                            'right': no_border,
                        }
                        if row_idx == 0:
                            borders['top'] = strong
                            borders['bottom'] = normal
                        if row_idx == row_count - 1:
                            borders['bottom'] = strong
                        _set_cell_borders(table.cell(row_idx, col_idx), **borders)

            def _docx_alignment(value):
                alignment = normalize_table_alignment(value)
                if alignment == TABLE_ALIGN_CENTER:
                    return WD_ALIGN_PARAGRAPH.CENTER
                if alignment == TABLE_ALIGN_RIGHT:
                    return WD_ALIGN_PARAGRAPH.RIGHT
                return WD_ALIGN_PARAGRAPH.LEFT

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = body_font_en
            style.font.size = Pt(body_pt)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

            # 添加标题
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    _set_run_font(run, h1_font, h1_font_en, h1_pt)

            def _write_body_paragraph(paragraph_text):
                p = doc.add_paragraph(paragraph_text)
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    _set_run_font(run, body_font, body_font_en, body_pt)

            def _write_table_block(block):
                rows = block.get('rows', [])
                if not isinstance(rows, list) or not rows:
                    return
                normalized_rows = []
                max_cols = 0
                for row in rows:
                    if isinstance(row, (list, tuple)):
                        normalized_row = [str(cell or '').replace('\r', '').replace('\n', ' ').strip() for cell in row]
                    else:
                        normalized_row = [str(row or '').strip()]
                    max_cols = max(max_cols, len(normalized_row))
                    normalized_rows.append(normalized_row)
                if not normalized_rows or max_cols <= 0:
                    return
                for row in normalized_rows:
                    if len(row) < max_cols:
                        row.extend([''] * (max_cols - len(row)))

                caption = str(block.get('caption', '') or '').strip()
                if caption:
                    caption_p = doc.add_paragraph(caption)
                    caption_p.paragraph_format.space_after = Pt(3)
                    for run in caption_p.runs:
                        _set_run_font(run, body_font, body_font_en, body_pt)

                table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
                if block.get('table_style') == TABLE_STYLE_THREE_LINE:
                    table.style = 'Normal Table'
                else:
                    table.style = 'Table Grid'
                merged_cells = normalize_merged_cells(
                    block.get('merged_cells', []),
                    len(normalized_rows),
                    max_cols,
                )
                cell_alignments = normalize_table_alignments(
                    block.get('cell_alignments', []),
                    len(normalized_rows),
                    max_cols,
                )
                covered_cells = set()
                merge_anchors = {(cell['row'], cell['col']): cell for cell in merged_cells}
                for merged in merged_cells:
                    for row_idx in range(merged['row'], merged['row'] + merged['rowspan']):
                        for col_idx in range(merged['col'], merged['col'] + merged['colspan']):
                            if row_idx == merged['row'] and col_idx == merged['col']:
                                continue
                            covered_cells.add((row_idx, col_idx))

                for row_idx, row in enumerate(normalized_rows):
                    for col_idx, cell_text in enumerate(row):
                        if (row_idx, col_idx) in covered_cells:
                            continue
                        cell = table.cell(row_idx, col_idx)
                        merged = merge_anchors.get((row_idx, col_idx))
                        if merged:
                            cell = cell.merge(
                                table.cell(
                                    row_idx + merged['rowspan'] - 1,
                                    col_idx + merged['colspan'] - 1,
                                )
                            )
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = _docx_alignment(cell_alignments[row_idx][col_idx])
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.first_line_indent = Pt(0)
                            for run in paragraph.runs:
                                _set_run_font(run, body_font, body_font_en, body_pt)
                if block.get('table_style') == TABLE_STYLE_THREE_LINE:
                    _apply_three_line_table_style(table)

            # 结构化导出（有章节数据时）
            if sections_data and isinstance(sections_data, dict):
                sd_order = sections_data.get('section_order', [])
                sd_sections = sections_data.get('sections', {})
                sd_levels = sections_data.get('section_levels', {})
                sd_blocks = sections_data.get('section_blocks', {})
                if sd_order and sd_sections:
                    level_fonts = {
                        1: (h1_font, h1_font_en, h1_pt),
                        2: (h2_font, h2_font_en, h2_pt),
                        3: (h3_font, h3_font_en, h3_pt),
                    }
                    for sec_title in sd_order:
                        sec_level = sd_levels.get(sec_title, 2)
                        if sec_title.strip():
                            heading_level = min(max(sec_level, 1), 3)
                            p = doc.add_heading(sec_title.strip(), level=heading_level)
                            cn, en, pt = level_fonts.get(heading_level, (h2_font, h2_font_en, h2_pt))
                            for run in p.runs:
                                _set_run_font(run, cn, en, pt)

                        sec_blocks = sd_blocks.get(sec_title, [])
                        if isinstance(sec_blocks, list) and sec_blocks:
                            for block in sanitize_blocks(sec_blocks):
                                if block['type'] == 'paragraph':
                                    if block['text']:
                                        _write_body_paragraph(block['text'])
                                    continue
                                if block['type'] == 'table':
                                    _write_table_block(block)
                            continue

                        sec_body = str(sd_sections.get(sec_title, '') or '').strip()
                        if sec_body:
                            for block in parse_markdown_blocks(sec_body):
                                if block['type'] == 'table':
                                    _write_table_block(block)
                                    continue
                                for para_text in str(block.get('text', '') or '').split('\n'):
                                    if para_text.strip():
                                        _write_body_paragraph(para_text)
                    doc.save(filepath)
                    return True

            # 按块添加内容（正则匹配模式，向后兼容）
            fallback_blocks = parse_markdown_blocks(text)
            if fallback_blocks:
                for block in fallback_blocks:
                    if block['type'] == 'table':
                        _write_table_block(block)
                        continue
                    for para_text in str(block.get('text', '') or '').split('\n'):
                        if not para_text.strip():
                            continue
                        if re.match(r'^第[一二三四五六七八九十\d]+[章]', para_text):
                            p = doc.add_heading(para_text, level=1)
                            for run in p.runs:
                                _set_run_font(run, h1_font, h1_font_en, h1_pt)
                        elif re.match(r'^[一二三四五六七八九十\d]+[、.．]', para_text) or \
                             re.match(r'^第[一二三四五六七八九十\d]+[节]', para_text) or \
                             re.match(r'^\d+\.\d+\s', para_text):
                            p = doc.add_heading(para_text, level=2)
                            for run in p.runs:
                                _set_run_font(run, h2_font, h2_font_en, h2_pt)
                        elif re.match(r'^\d+\.\d+\.\d+\s', para_text) or \
                             re.match(r'^（[一二三四五六七八九十\d]+）', para_text):
                            p = doc.add_heading(para_text, level=3)
                            for run in p.runs:
                                _set_run_font(run, h3_font, h3_font_en, h3_pt)
                        else:
                            _write_body_paragraph(para_text)
            else:
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if not para_text.strip():
                        continue
                    if re.match(r'^第[一二三四五六七八九十\d]+[章]', para_text):
                        p = doc.add_heading(para_text, level=1)
                        for run in p.runs:
                            _set_run_font(run, h1_font, h1_font_en, h1_pt)
                    elif re.match(r'^[一二三四五六七八九十\d]+[、.．]', para_text) or \
                         re.match(r'^第[一二三四五六七八九十\d]+[节]', para_text) or \
                         re.match(r'^\d+\.\d+\s', para_text):
                        p = doc.add_heading(para_text, level=2)
                        for run in p.runs:
                            _set_run_font(run, h2_font, h2_font_en, h2_pt)
                    elif re.match(r'^\d+\.\d+\.\d+\s', para_text) or \
                         re.match(r'^（[一二三四五六七八九十\d]+）', para_text):
                        p = doc.add_heading(para_text, level=3)
                        for run in p.runs:
                            _set_run_font(run, h3_font, h3_font_en, h3_pt)
                    else:
                        _write_body_paragraph(para_text)

            doc.save(filepath)
            return True
        except ImportError:
            raise RuntimeError('请安装python-docx库')
        except Exception as e:
            raise RuntimeError(f'导出Word失败: {e}')

    def export_doc(self, text: str, filepath: str, title: str = '', level_font_styles: dict = None, sections_data: dict = None) -> bool:
        """导出为 DOC 文档（依赖本机 Microsoft Word）"""
        tmp_docx = ''
        word = None
        document = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_docx = tmp_file.name

            self.export_docx(text, tmp_docx, title, level_font_styles=level_font_styles, sections_data=sections_data)

            try:
                import win32com.client  # type: ignore
            except Exception as exc:
                raise RuntimeError('导出 DOC 需要本机安装 Microsoft Word') from exc

            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(os.path.abspath(tmp_docx), ReadOnly=False, AddToRecentFiles=False)
            try:
                document.SaveAs2(os.path.abspath(filepath), FileFormat=0)
            except Exception:
                document.SaveAs(os.path.abspath(filepath), FileFormat=0)
            return True
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f'导出DOC失败: {e}')
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            if tmp_docx and os.path.exists(tmp_docx):
                try:
                    os.remove(tmp_docx)
                except OSError:
                    pass

    def export_latex(self, text: str, filepath: str, title: str = '') -> bool:
        """导出为 LaTeX 源文件。"""
        try:
            safe_title = self._escape_latex(str(title or '').replace('\n', ' ').strip())
            body = self._latex_body_from_text(text)
            document_lines = [
                '% !TEX program = xelatex',
                r'\documentclass[UTF8]{ctexart}',
                r'\usepackage[a4paper,margin=2.54cm]{geometry}',
                r'\usepackage{setspace}',
                r'\setstretch{1.35}',
                '',
            ]

            if safe_title:
                document_lines.extend(
                    [
                        rf'\title{{{safe_title}}}',
                        r'\author{}',
                        r'\date{}',
                        '',
                    ]
                )

            document_lines.append(r'\begin{document}')
            if safe_title:
                document_lines.extend([r'\maketitle', ''])
            if body:
                document_lines.append(body)
            document_lines.extend(['', r'\end{document}', ''])

            with open(filepath, 'w', encoding='utf-8', newline='\n') as f:
                f.write('\n'.join(document_lines))
            return True
        except Exception as e:
            raise RuntimeError(f'导出LaTeX失败: {e}')

    def export_txt(self, text: str, filepath: str) -> bool:
        """导出为TXT文件"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            raise RuntimeError(f'导出TXT失败: {e}')

    def export_pdf(self, text: str, filepath: str, title: str = '') -> bool:
        """导出为PDF（通过Word或LibreOffice转换）"""
        # 先导出为docx，再尝试转PDF
        base = os.path.splitext(filepath)[0]
        tmp_docx = base + '_tmp.docx'
        try:
            self.export_docx(text, tmp_docx, title)
            # Windows 优先使用 Word COM 自动化
            if sys.platform == 'win32':
                word = None
                doc = None
                try:
                    import comtypes.client
                    word = comtypes.client.CreateObject('Word.Application')
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(tmp_docx))
                    doc.SaveAs(os.path.abspath(filepath), FileFormat=17)  # wdFormatPDF = 17
                    return True
                except Exception:
                    pass
                finally:
                    if doc is not None:
                        try:
                            doc.Close(False)
                        except Exception:
                            pass
                    if word is not None:
                        try:
                            word.Quit()
                        except Exception:
                            pass
            # 回退到 LibreOffice
            import subprocess
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(filepath), tmp_docx],
                capture_output=True, timeout=30
            )
            if result.returncode == 0:
                return True
            raise RuntimeError('PDF 转换工具不可用，请手动从 Word 另存为 PDF')
        except FileNotFoundError:
            raise RuntimeError('未找到 PDF 转换工具，请从 Word 文件另存为 PDF')
        finally:
            if os.path.exists(tmp_docx):
                os.remove(tmp_docx)

    def diff_text(self, text1: str, text2: str) -> list:
        """对比两段文本，返回差异列表"""
        lines1 = text1.splitlines(keepends=True)
        lines2 = text2.splitlines(keepends=True)
        differ = difflib.unified_diff(
            lines1, lines2,
            fromfile='原文', tofile='修改后',
            lineterm=''
        )
        return list(differ)

    def diff_highlight(self, text1: str, text2: str) -> list:
        """返回带标记的差异列表，用于GUI高亮显示"""
        matcher = difflib.SequenceMatcher(None, text1, text2)
        result = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                result.append(('equal', text1[i1:i2]))
            elif tag == 'insert':
                result.append(('insert', text2[j1:j2]))
            elif tag == 'delete':
                result.append(('delete', text1[i1:i2]))
            elif tag == 'replace':
                result.append(('delete', text1[i1:i2]))
                result.append(('insert', text2[j1:j2]))
        return result

    def count_words(self, text: str) -> dict:
        """统计字数信息"""
        cn_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        en_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        numbers = len(re.findall(r'\b\d+\b', text))
        sentences = len(re.findall(r'[。！？.!?]+', text))
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        return {
            'total': len(text),
            'chinese': cn_chars,
            'english_words': en_words,
            'numbers': numbers,
            'sentences': sentences,
            'paragraphs': paragraphs,
        }

    def check_format(self, text: str, style: str = '学术论文') -> dict:
        """执行本地格式与结构检查。"""
        content = str(text or '')
        issues = []

        if ',' in content and '，' not in content:
            issues.append('建议使用中文逗号（，）替代英文逗号（,）')
        if '.' in content.replace('...', '') and '。' not in content:
            issues.append('建议使用中文句号（。）替代英文句号（.）')

        cn_nums = re.findall(r'[一二三四五六七八九十百千万]+', content)
        if cn_nums:
            issues.append(f'发现{len(cn_nums)}处中文数字，{style}建议优先使用阿拉伯数字')

        paragraphs = [p for p in content.split('\n') if p.strip()]
        short_paras = [p for p in paragraphs if len(p.strip()) < 50]
        if short_paras:
            issues.append(f'发现{len(short_paras)}个过短段落，建议合并或扩充')

        has_ref = bool(re.search(r'\[\d+\]', content))
        if len(content) > 500 and not has_ref:
            issues.append('未发现参考文献引用标注，建议补充引用编号')

        sentence_count = len([s for s in re.split(r'[。！？.!?]', content) if s.strip()])
        return {
            'issues': issues,
            'issue_count': len(issues),
            'word_count': len(content),
            'para_count': len(paragraphs),
            'sentence_count': sentence_count,
        }

    def detect_sensitive(self, text: str) -> list:
        """检测敏感内容"""
        sensitive_patterns = [
            (r'作弊|抄袭|代写|枪手', '学术诚信风险'),
            (r'政治|党|政府|领导人', '政治敏感内容'),
            (r'色情|暴力|恐怖', '违规内容'),
        ]
        found = []
        for pattern, category in sensitive_patterns:
            matches = re.findall(pattern, text)
            if matches:
                found.append({'category': category, 'matches': matches[:5]})
        return found
